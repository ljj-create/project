"""
文档加载器 — 支持 PDF、Markdown、纯文本、代码文件
"""
from pathlib import Path
from dataclasses import dataclass
from loguru import logger


@dataclass
class Document:
    """文档单元"""
    content: str
    metadata: dict
    source: str  # 文件路径或 URL


class DocumentLoader:
    """
    多格式文档加载器

    支持格式:
    - PDF (.pdf)
    - Markdown (.md)
    - 纯文本 (.txt)
    - 代码文件 (.py, .java, .cpp, .c, .go, .js, .ts)
    - Word (.docx)
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".md", ".txt", ".markdown",
        ".py", ".java", ".cpp", ".c", ".h", ".go", ".js", ".ts", ".rs",
        ".json", ".yaml", ".yml", ".toml",
        ".docx",
    }

    @classmethod
    def load_file(cls, file_path: str | Path) -> Document:
        """加载单个文件"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        ext = path.suffix.lower()
        if ext not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        logger.info(f"加载文件: {path}")

        if ext == ".pdf":
            return cls._load_pdf(path)
        elif ext in (".md", ".markdown"):
            return cls._load_markdown(path)
        elif ext == ".docx":
            return cls._load_docx(path)
        else:
            return cls._load_text(path)

    @classmethod
    def load_directory(cls, dir_path: str | Path, recursive: bool = True) -> list[Document]:
        """加载目录下的所有支持文件"""
        path = Path(dir_path)
        if not path.is_dir():
            raise NotADirectoryError(f"目录不存在: {path}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in sorted(path.glob(pattern)):
            if file_path.is_file() and file_path.suffix.lower() in cls.SUPPORTED_EXTENSIONS:
                try:
                    doc = cls.load_file(file_path)
                    documents.append(doc)
                except Exception as e:
                    logger.warning(f"跳过文件 {file_path}: {e}")

        logger.info(f"从 {path} 加载了 {len(documents)} 个文件")
        return documents

    @classmethod
    def _load_pdf(cls, path: Path) -> Document:
        """加载 PDF 文件"""
        pages = 0
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            pages = len(reader.pages)
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            content = "\n\n".join(text_parts)
        except ImportError:
            logger.warning("PyPDF2 未安装，尝试使用 unstructured")
            from unstructured.partition.pdf import partition_pdf
            elements = partition_pdf(str(path))
            content = "\n\n".join(str(el) for el in elements)

        return Document(
            content=content,
            metadata={"source": str(path), "type": "pdf", "pages": pages},
            source=str(path),
        )

    @classmethod
    def _load_markdown(cls, path: Path) -> Document:
        """加载 Markdown 文件"""
        content = path.read_text(encoding="utf-8")
        return Document(
            content=content,
            metadata={"source": str(path), "type": "markdown"},
            source=str(path),
        )

    @classmethod
    def _load_docx(cls, path: Path) -> Document:
        """加载 Word 文档"""
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)
        return Document(
            content=content,
            metadata={"source": str(path), "type": "docx"},
            source=str(path),
        )

    @classmethod
    def _load_text(cls, path: Path) -> Document:
        """加载纯文本/代码文件"""
        try:
            content = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            content = path.read_text(encoding="utf-8", errors="replace")
        # 为代码文件添加语言标记
        lang_map = {
            ".py": "python", ".java": "java", ".cpp": "cpp", ".c": "c",
            ".h": "c", ".go": "go", ".js": "javascript", ".ts": "typescript",
            ".rs": "rust", ".json": "json", ".yaml": "yaml", ".yml": "yaml",
            ".toml": "toml",
        }
        lang = lang_map.get(path.suffix.lower(), "text")
        if lang in ("python", "java", "cpp", "c", "go", "javascript", "typescript", "rust"):
            content = f"```{lang}\n{content}\n```"

        return Document(
            content=content,
            metadata={"source": str(path), "type": "code", "language": lang},
            source=str(path),
        )
